"""File format handlers for different document types."""

from __future__ import annotations

import json
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Callable, Optional

# Optional imports - will be None if not installed
try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from odf import text as odf_text
    from odf.opendocument import load as odf_load
except ImportError:
    odf_text = None
    odf_load = None


# File extensions grouped by handler type
PLAIN_TEXT_EXTENSIONS = {
    # Plain text
    ".txt", ".text", ".log", ".md", ".rst", ".asc",
    # Web/markup
    ".html", ".htm", ".xhtml", ".xml", ".csv", ".tsv",
    ".json", ".yaml", ".yml",
    # Documentation
    ".tex", ".latex", ".adoc", ".org", ".wiki", ".mediawiki",
    # Config files
    ".ini", ".cfg", ".conf", ".properties", ".env", ".toml", ".lock",
    # Source code
    ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".cs",
    ".php", ".rb", ".go", ".rs", ".sh", ".bat", ".ps1", ".sql",
    # Notes/misc
    ".note", ".eml", ".mbox",
    # Subtitles
    ".srt", ".vtt", ".sub",
    # Translation
    ".po", ".pot",
    # Other
    ".texi", ".man", ".nfo", ".readme",
}

DOCX_EXTENSIONS = {".docx", ".docm"}
XLSX_EXTENSIONS = {".xlsx", ".xlsm"}
ODT_EXTENSIONS = {".odt"}
RTF_EXTENSIONS = {".rtf"}


class FileHandler(ABC):
    """Base class for file format handlers."""

    @abstractmethod
    def read(self, path: Path) -> str:
        """Read file and return text content."""
        pass

    @abstractmethod
    def convert(
        self,
        input_path: Path,
        output_path: Path,
        converter: Callable[[str], str],
    ) -> None:
        """Convert file in-place or to new file, preserving format."""
        pass

    @staticmethod
    def can_handle(path: Path) -> bool:
        """Check if this handler can process the given file."""
        return False


class PlainTextHandler(FileHandler):
    """Handler for plain text files."""

    def read(self, path: Path, encoding: str = "utf-8") -> str:
        return path.read_text(encoding=encoding)

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        converter: Callable[[str], str],
        encoding: str = "utf-8",
    ) -> None:
        text = self.read(input_path, encoding=encoding)
        converted = converter(text)
        output_path.write_text(converted, encoding=encoding)

    @staticmethod
    def can_handle(path: Path) -> bool:
        return path.suffix.lower() in PLAIN_TEXT_EXTENSIONS or path.suffix == ""


class DocxHandler(FileHandler):
    """Handler for Microsoft Word .docx files."""

    def __init__(self):
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required for .docx support. "
                "Install with: pip install paraencoder[office]"
            )

    def read(self, path: Path) -> str:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        converter: Callable[[str], str],
    ) -> None:
        doc = DocxDocument(str(input_path))

        # Convert text in paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text:
                    run.text = converter(run.text)

        # Convert text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                run.text = converter(run.text)

        # Convert text in headers/footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        for run in para.runs:
                            if run.text:
                                run.text = converter(run.text)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        for run in para.runs:
                            if run.text:
                                run.text = converter(run.text)

        doc.save(str(output_path))

    @staticmethod
    def can_handle(path: Path) -> bool:
        return path.suffix.lower() in DOCX_EXTENSIONS


class XlsxHandler(FileHandler):
    """Handler for Microsoft Excel .xlsx files."""

    def __init__(self):
        if openpyxl is None:
            raise ImportError(
                "openpyxl is required for .xlsx support. "
                "Install with: pip install paraencoder[office]"
            )

    def read(self, path: Path) -> str:
        wb = openpyxl.load_workbook(str(path), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        lines.append(cell.value)
        return "\n".join(lines)

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        converter: Callable[[str], str],
    ) -> None:
        # Load workbook preserving everything (images, charts, etc.)
        wb = openpyxl.load_workbook(str(input_path))

        for sheet in wb.worksheets:
            # Convert regular cell values
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        cell.value = converter(cell.value)

            # Convert merged cell values (they're stored in the top-left cell)
            for merged_range in sheet.merged_cells.ranges:
                cell = sheet.cell(merged_range.min_row, merged_range.min_col)
                if cell.value and isinstance(cell.value, str):
                    cell.value = converter(cell.value)

            # Convert comments
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.comment and cell.comment.text:
                        cell.comment.text = converter(cell.comment.text)

            # Convert header/footer
            if sheet.oddHeader and sheet.oddHeader.center:
                if sheet.oddHeader.center.text:
                    sheet.oddHeader.center.text = converter(sheet.oddHeader.center.text)
            if sheet.oddHeader and sheet.oddHeader.left:
                if sheet.oddHeader.left.text:
                    sheet.oddHeader.left.text = converter(sheet.oddHeader.left.text)
            if sheet.oddHeader and sheet.oddHeader.right:
                if sheet.oddHeader.right.text:
                    sheet.oddHeader.right.text = converter(sheet.oddHeader.right.text)
            if sheet.oddFooter and sheet.oddFooter.center:
                if sheet.oddFooter.center.text:
                    sheet.oddFooter.center.text = converter(sheet.oddFooter.center.text)

        # Convert sheet names
        for sheet in wb.worksheets:
            original_title = sheet.title
            converted_title = converter(original_title)
            if converted_title != original_title:
                sheet.title = converted_title

        wb.save(str(output_path))

    @staticmethod
    def can_handle(path: Path) -> bool:
        return path.suffix.lower() in XLSX_EXTENSIONS


class OdtHandler(FileHandler):
    """Handler for OpenDocument .odt files."""

    def __init__(self):
        if odf_load is None:
            raise ImportError(
                "odfpy is required for .odt support. "
                "Install with: pip install paraencoder[office]"
            )

    def _get_text_elements(self, element):
        """Recursively get all text elements."""
        from odf.text import P, H, Span
        from odf.element import Text

        elements = []
        for child in element.childNodes:
            if isinstance(child, (P, H, Span)):
                elements.extend(self._get_text_elements(child))
            elif isinstance(child, Text):
                elements.append(child)
            elif hasattr(child, 'childNodes'):
                elements.extend(self._get_text_elements(child))
        return elements

    def read(self, path: Path) -> str:
        doc = odf_load(str(path))
        text_content = []
        for para in doc.getElementsByType(odf_text.P):
            text_content.append(str(para))
        return "\n".join(text_content)

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        converter: Callable[[str], str],
    ) -> None:
        doc = odf_load(str(input_path))

        # Get all text elements and convert them
        text_elements = self._get_text_elements(doc.body)
        for text_node in text_elements:
            if text_node.data:
                text_node.data = converter(text_node.data)

        doc.save(str(output_path))

    @staticmethod
    def can_handle(path: Path) -> bool:
        return path.suffix.lower() in ODT_EXTENSIONS


def get_handler(path: Path) -> FileHandler:
    """Get the appropriate handler for a file path."""
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in DOCX_EXTENSIONS:
        return DocxHandler()
    elif suffix in XLSX_EXTENSIONS:
        return XlsxHandler()
    elif suffix in ODT_EXTENSIONS:
        return OdtHandler()
    elif suffix in PLAIN_TEXT_EXTENSIONS or suffix == "":
        return PlainTextHandler()
    else:
        # Try plain text as fallback
        return PlainTextHandler()


def get_supported_extensions() -> set[str]:
    """Get all supported file extensions."""
    extensions = set(PLAIN_TEXT_EXTENSIONS)
    extensions.update(DOCX_EXTENSIONS)
    extensions.update(XLSX_EXTENSIONS)
    extensions.update(ODT_EXTENSIONS)
    return extensions


def is_supported(path: Path) -> bool:
    """Check if a file type is supported."""
    path = Path(path)
    suffix = path.suffix.lower()
    return suffix in get_supported_extensions() or suffix == ""
